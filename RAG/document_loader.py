# ==================== DOCUMENT LOADER ====================
"""
Document Loader for BPJS RAG System
Supports PDF, DOCX, and TXT files
"""

import os
from pathlib import Path
from typing import Tuple, Dict, List
from RAG.rag_config import RAGConfig
import PyPDF2
from docx import Document as DocxDocument


class DocumentLoader:
    """Load and preprocess documents from various formats"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
    
    def load_pdf(self, file_path: str) -> str:
        """
        Load PDF document and extract text
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        print(f"  📄 Loading PDF: {Path(file_path).name}")
        
        if not os.path.exists(file_path):
            print(f"  ⚠️  File not found: {file_path}")
            return f"[PLACEHOLDER] Content from {Path(file_path).name}"
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                print(f"     Pages: {num_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text
                    
                    # Progress indicator for large PDFs
                    if (page_num + 1) % 10 == 0:
                        print(f"     Extracted: {page_num + 1}/{num_pages} pages")
                
                # Clean up common PDF extraction artifacts
                text = self._clean_pdf_text(text)
                
                print(f"  ✓ Extracted {len(text):,} characters")
                return text
                
        except Exception as e:
            print(f"  ✗ Error loading PDF: {e}")
            return f"[ERROR] Could not load {Path(file_path).name}: {str(e)}"
    
    def load_docx(self, file_path: str) -> str:
        """
        Load DOCX document and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        print(f"  📄 Loading DOCX: {Path(file_path).name}")
        
        if not os.path.exists(file_path):
            print(f"  ⚠️  File not found: {file_path}")
            return f"[PLACEHOLDER] Content from {Path(file_path).name}"
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            print(f"  ✓ Extracted {len(text):,} characters")
            return text
            
        except Exception as e:
            print(f"  ✗ Error loading DOCX: {e}")
            return f"[ERROR] Could not load {Path(file_path).name}: {str(e)}"
    
    def load_txt(self, file_path: str, encoding: str = 'utf-8') -> str:
        """
        Load plain text file
        
        Args:
            file_path: Path to text file
            encoding: Text encoding (default: utf-8)
            
        Returns:
            Text content
        """
        print(f"  📄 Loading TXT: {Path(file_path).name}")
        
        if not os.path.exists(file_path):
            print(f"  ⚠️  File not found: {file_path}")
            return f"[PLACEHOLDER] Content from {Path(file_path).name}"
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            
            print(f"  ✓ Loaded {len(text):,} characters")
            return text
            
        except UnicodeDecodeError:
            # Try alternative encoding
            print(f"  ⚠️  UTF-8 failed, trying latin-1...")
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                print(f"  ✓ Loaded {len(text):,} characters (latin-1)")
                return text
            except Exception as e:
                print(f"  ✗ Error: {e}")
                return f"[ERROR] Could not decode {Path(file_path).name}"
        
        except Exception as e:
            print(f"  ✗ Error loading TXT: {e}")
            return f"[ERROR] Could not load {Path(file_path).name}: {str(e)}"
    
    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean common PDF extraction artifacts
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove common PDF artifacts
        text = text.replace('\x00', '')  # Null bytes
        text = text.replace('\uf0b7', '•')  # Bullet points
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n')
        text = text.replace('\r', '\n')
        
        # Remove duplicate spaces
        import re
        text = re.sub(r' +', ' ', text)
        
        return text
    
    def load_document(self, doc_key: str) -> Tuple[str, Dict]:
        """
        Load document by key and return text + metadata
        
        Args:
            doc_key: Document key from RAGConfig.DOCUMENT_PATHS
            
        Returns:
            Tuple of (text_content, metadata_dict)
        """
        file_path = self.config.DOCUMENT_PATHS.get(doc_key)
        
        if not file_path:
            print(f"  ✗ Unknown document key: {doc_key}")
            return f"[ERROR] Unknown document: {doc_key}", {}
        
        metadata = self.config.DOCUMENT_METADATA.get(doc_key, {})
        
        # Determine file type and load
        file_path_lower = file_path.lower()
        
        if file_path_lower.endswith('.pdf'):
            text = self.load_pdf(file_path)
        elif file_path_lower.endswith('.docx'):
            text = self.load_docx(file_path)
        elif file_path_lower.endswith('.txt'):
            text = self.load_txt(file_path)
        else:
            print(f"  ⚠️  Unsupported format: {file_path}")
            text = f"[PLACEHOLDER] Unsupported format: {Path(file_path).name}"
        
        # Add file info to metadata
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            metadata['file_path'] = file_path
            metadata['file_size_bytes'] = file_size
            metadata['file_size_mb'] = round(file_size / (1024 * 1024), 2)
        
        return text, metadata
    
    def load_all_documents(self) -> List[Tuple[str, Dict]]:
        """
        Load all configured documents
        
        Returns:
            List of (text, metadata) tuples
        """
        print("\n" + "="*80)
        print("LOADING DOCUMENTS")
        print("="*80)
        
        documents = []
        
        for doc_key in self.config.DOCUMENT_PATHS.keys():
            print(f"\n[{doc_key}]")
            text, metadata = self.load_document(doc_key)
            documents.append((text, metadata))
            
            # Summary
            if text.startswith("[ERROR]") or text.startswith("[PLACEHOLDER]"):
                print(f"  ⚠️  Failed to load document")
            else:
                print(f"  ✓ Loaded: {len(text):,} characters")
                if 'file_size_mb' in metadata:
                    print(f"  📊 File size: {metadata['file_size_mb']:.2f} MB")
        
        print(f"\n{'='*80}")
        print(f"✅ DOCUMENTS LOADED: {len(documents)}")
        
        # Calculate total size
        total_chars = sum(len(text) for text, _ in documents if not text.startswith('['))
        total_tokens_estimate = total_chars // 4
        
        print(f"   Total characters: {total_chars:,}")
        print(f"   Estimated tokens: {total_tokens_estimate:,}")
        print(f"{'='*80}\n")
        
        return documents


# ==================== HELPER FUNCTION ====================

def validate_document_paths(config: RAGConfig) -> Dict[str, bool]:
    """
    Validate that all configured document paths exist
    
    Args:
        config: RAGConfig instance
        
    Returns:
        Dict mapping doc_key to exists status
    """
    print("\n" + "="*80)
    print("VALIDATING DOCUMENT PATHS")
    print("="*80)
    
    results = {}
    
    for doc_key, file_path in config.DOCUMENT_PATHS.items():
        exists = os.path.exists(file_path)
        results[doc_key] = exists
        
        status = "✓" if exists else "✗"
        print(f"{status} {doc_key}")
        print(f"  {file_path}")
        
        if exists:
            file_size = os.path.getsize(file_path)
            print(f"  {file_size:,} bytes ({file_size/(1024*1024):.2f} MB)")
    
    total = len(results)
    found = sum(results.values())
    
    print(f"\n{'='*80}")
    print(f"VALIDATION SUMMARY: {found}/{total} documents found")
    print(f"{'='*80}\n")
    
    return results


# ==================== CLI TESTING ====================

if __name__ == "__main__":
    """Test document loader standalone"""
    from rag_config import RAGConfig
    
    config = RAGConfig()
    
    # Validate paths first
    validation = validate_document_paths(config)
    
    if not any(validation.values()):
        print("⚠️  No documents found. Check paths in rag_config.py")
        exit(1)
    
    # Load documents
    loader = DocumentLoader(config)
    documents = loader.load_all_documents()
    
    # Show sample from first document
    if documents:
        text, metadata = documents[0]
        print("\n" + "="*80)
        print("SAMPLE FROM FIRST DOCUMENT")
        print("="*80)
        print(f"Title: {metadata.get('title', 'Unknown')}")
        print(f"Type: {metadata.get('document_type', 'Unknown')}")
        print(f"\nFirst 500 characters:")
        print("-"*80)
        print(text[:500])
        print("...")
